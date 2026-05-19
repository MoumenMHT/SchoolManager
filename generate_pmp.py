#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de Plan de Management de Projet (PMP) - SchoolHub
Ce script crée un document Word professionnel pour le PMP du projet SchoolHub
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

# Configuration de styles
def setup_styles(doc):
    """Configure les styles du document"""
    styles = doc.styles
    
    # Style pour les titres principaux
    heading1_style = styles['Heading 1']
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Style pour les sous-titres
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(13)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 102, 153)

def add_title_page(doc):
    """Ajoute la page de titre"""
    # Titre principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLAN DE MANAGEMENT DE PROJET")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SchoolHub - Système de Gestion d'Établissement Scolaire")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    # Espace
    doc.add_paragraph()
    
    # Informations du projet
    info_section = doc.add_paragraph()
    info_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_text = f"""
Projet de Fin d'Études (PFE)
Présenté le : {datetime.now().strftime('%d/%m/%Y')}
Remise prévue : 21/05/2026
Version : 1.0
    """
    
    for line in info_text.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    # Page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Ajoute une table des matières"""
    doc.add_heading("Table des Matières", level=1)
    
    contents = [
        "1. Contexte, Enjeux et Objectifs du Projet",
        "2. Organigramme des Tâches (WBS)",
        "3. Organigramme des Parties Prenantes",
        "4. Planning Prévisionnel et Échéances",
        "5. Indicateurs de Suivi",
        "6. Livrables Identifiés",
        "7. Budget Prévisionnel",
        "8. Ressources Techniques et Humaines",
        "9. Analyse des Risques",
        "10. Stratégie de Communication",
    ]
    
    for item in contents:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()

def add_context_section(doc):
    """Ajoute la section contexte et objectifs"""
    doc.add_heading("1. Contexte, Enjeux et Objectifs du Projet", level=1)
    
    doc.add_heading("1.1 Contexte du Projet", level=2)
    doc.add_paragraph(
        "SchoolHub est un système complet de gestion d'établissements scolaires développé dans le cadre "
        "d'un projet de fin d'études. Le projet vise à moderniser et automatiser les processus administratifs "
        "et académiques des institutions éducatives."
    )
    
    doc.add_heading("1.2 Enjeux Majeurs", level=2)
    enjeux = [
        "Améliorer la gestion administrative des établissements scolaires",
        "Fournir un accès centralisé aux données étudiantes et pédagogiques",
        "Optimiser la communication entre administrateurs, enseignants et parents",
        "Automatiser les processus de gestion (notes, présences, paiements)",
        "Offrir une solution scalable et maintenable"
    ]
    for enjeu in enjeux:
        doc.add_paragraph(enjeu, style='List Bullet')
    
    doc.add_heading("1.3 Objectifs du Projet", level=2)
    doc.add_heading("Objectifs Primaires", level=3)
    objectifs = [
        "Développer une API REST robuste et sécurisée pour la gestion d'école",
        "Créer une interface web intuitive et responsive pour les utilisateurs",
        "Implémenter un système d'authentification multi-rôles (Admin, Professeur, Parent)",
        "Gérer complètement les données académiques (étudiants, notes, présences, emplois du temps)",
        "Intégrer un système de gestion des paiements et finances"
    ]
    for obj in objectifs:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading("Objectifs Secondaires", level=3)
    obj_sec = [
        "Fournir des dashboards et rapports analytiques",
        "Assurer une haute disponibilité du système",
        "Documenter le projet de manière complète",
        "Implémenter des tests unitaires et d'intégration",
        "Optimiser les performances du système"
    ]
    for obj in obj_sec:
        doc.add_paragraph(obj, style='List Bullet')

def add_wbs_section(doc):
    """Ajoute la section WBS (Organigramme des tâches)"""
    doc.add_heading("2. Organigramme des Tâches (WBS)", level=1)
    
    doc.add_paragraph(
        "Le projet est décomposé en workstreams principaux, chacun contenant des sous-tâches."
    )
    
    wbs_data = {
        "SchoolHub": {
            "Phase 1 - Planification & Design (Semaine 1-2)": [
                "Analyse des besoins et spécifications",
                "Conception de l'architecture système",
                "Design de la base de données",
                "Maquettes de l'interface utilisateur",
                "Planification détaillée"
            ],
            "Phase 2 - Backend (Semaine 3-5)": [
                "Configuration de l'environnement Laravel",
                "Implémentation de l'authentification",
                "Développement des modèles Eloquent",
                "Création des migrations de base de données",
                "Implémentation des controllers API",
                "Gestion des autorisations et rôles",
                "Tests unitaires backend"
            ],
            "Phase 3 - Frontend (Semaine 3-5)": [
                "Configuration de l'environnement Vue.js",
                "Développement des composants réutilisables",
                "Création des pages de gestion",
                "Intégration avec l'API",
                "Responsive design et UX",
                "Tests unitaires frontend"
            ],
            "Phase 4 - Intégration & Testing (Semaine 6-7)": [
                "Tests d'intégration système",
                "Tests de performance et optimisation",
                "Tests de sécurité",
                "Déploiement sur serveur de test",
                "Correction des bugs majeurs"
            ],
            "Phase 5 - Finalisation & Documentation (Semaine 8)": [
                "Rédaction de la documentation utilisateur",
                "Rédaction de la documentation technique",
                "Création de tutoriels et guides",
                "Préparation de la démonstration",
                "Déploiement en production"
            ]
        }
    }
    
    for phase, sub_items in wbs_data["SchoolHub"].items():
        p = doc.add_paragraph(phase, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        for item in sub_items:
            sub_p = doc.add_paragraph(item, style='List Bullet 2')
            sub_p.paragraph_format.left_indent = Inches(0.6)

def add_stakeholders_section(doc):
    """Ajoute la section des parties prenantes"""
    doc.add_heading("3. Organigramme des Parties Prenantes", level=1)
    
    doc.add_paragraph("Identification et rôles des parties prenantes du projet :")
    
    # Créer un tableau des parties prenantes
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # En-têtes du tableau
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Partie Prenante"
    hdr_cells[1].text = "Rôle"
    hdr_cells[2].text = "Responsabilités"
    hdr_cells[3].text = "Intérêt"
    
    # Données
    stakeholders = [
        ("Chef de Projet", "Pilote", "Gestion globale, coordination, reporting", "Succès du projet"),
        ("Équipe Développement Backend", "Exécutant", "Développement API, base de données", "Qualité du code"),
        ("Équipe Développement Frontend", "Exécutant", "Interface utilisateur, intégration", "Expérience utilisateur"),
        ("Sponsor/Encadrant", "Décideur", "Orientation, validation, clôture", "Atteinte des objectifs"),
        ("Utilisateurs Finaux", "Bénéficiaire", "Tests, feedback, utilisation", "Satisfaction des besoins"),
    ]
    
    for stakeholder, role, resp, interest in stakeholders:
        row_cells = table.add_row().cells
        row_cells[0].text = stakeholder
        row_cells[1].text = role
        row_cells[2].text = resp
        row_cells[3].text = interest

def add_schedule_section(doc):
    """Ajoute la section du planning"""
    doc.add_heading("4. Planning Prévisionnel et Échéances", level=1)
    
    doc.add_heading("4.1 Phases du Projet", level=2)
    
    # Planning détaillé
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Phase"
    hdr_cells[1].text = "Début"
    hdr_cells[2].text = "Fin"
    hdr_cells[3].text = "Durée"
    hdr_cells[4].text = "Jalon"
    
    phases = [
        ("Planification & Design", "08/05/2026", "12/05/2026", "5 jours", "Design approuvé"),
        ("Développement Backend", "13/05/2026", "17/05/2026", "5 jours", "API fonctionnelle"),
        ("Développement Frontend", "13/05/2026", "17/05/2026", "5 jours", "UI complète"),
        ("Intégration & Testing", "18/05/2026", "20/05/2026", "3 jours", "Tests validés"),
        ("Documentation & Finalisation", "21/05/2026", "21/05/2026", "1 jour", "PMP remis"),
    ]
    
    for phase, debut, fin, duree, jalon in phases:
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = debut
        row_cells[2].text = fin
        row_cells[3].text = duree
        row_cells[4].text = jalon
    
    doc.add_heading("4.2 Jalons Critiques", level=2)
    jalons = [
        ("21/05/2026", "Remise du Plan de Management de Projet (PMP)"),
        ("20/05/2026", "Finalisation des tests et validation système"),
        ("17/05/2026", "Fin du développement backend et frontend"),
        ("25/06/2026", "Remise du rapport de stage"),
    ]
    
    for date, jalon in jalons:
        doc.add_paragraph(f"{date} : {jalon}", style='List Bullet')

def add_indicators_section(doc):
    """Ajoute la section des indicateurs de suivi"""
    doc.add_heading("5. Indicateurs de Suivi", level=1)
    
    doc.add_heading("5.1 Indicateurs de Performance", level=2)
    
    indicators = {
        "Complétude": "% des tâches planifiées complétées à temps",
        "Qualité": "Nombre de bugs par 100 lignes de code",
        "Respect du Planning": "% des jalons respectés",
        "Ressources": "% de l'effort prévu utilisé",
        "Test Coverage": "% du code couvert par les tests",
    }
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Indicateur"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Cible"
    
    for indicator, description in indicators.items():
        row_cells = table.add_row().cells
        row_cells[0].text = indicator
        row_cells[1].text = description
        if indicator == "Complétude":
            row_cells[2].text = ">= 95%"
        elif indicator == "Respect du Planning":
            row_cells[2].text = "100%"
        else:
            row_cells[2].text = "À définir"

def add_deliverables_section(doc):
    """Ajoute la section des livrables"""
    doc.add_heading("6. Livrables Identifiés", level=1)
    
    deliverables = {
        "Livrables Techniques": [
            "Code source du backend (API Laravel)",
            "Code source du frontend (Interface Vue.js)",
            "Base de données relationnelle configurée",
            "Documentation API complète (OpenAPI/Swagger)",
            "Guide d'installation et de déploiement",
            "Suite de tests automatisés",
            "Scripts de migration de données"
        ],
        "Livrables de Documentation": [
            "Plan de Management de Projet (PMP)",
            "Rapport de fin d'études",
            "Architecture système et design patterns",
            "Manuel utilisateur",
            "Guide administrateur",
            "Documentation du code source"
        ],
        "Livrables de Démonstration": [
            "Démonstration fonctionnelle du système",
            "Vidéo tutoriel d'utilisation",
            "Présentation des résultats"
        ]
    }
    
    for category, items in deliverables.items():
        doc.add_heading(category, level=2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

def add_budget_section(doc):
    """Ajoute la section du budget"""
    doc.add_heading("7. Budget Prévisionnel", level=1)
    
    doc.add_heading("7.1 Ressources Informatiques", level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Élément"
    hdr_cells[1].text = "Quantité"
    hdr_cells[2].text = "Coût Unitaire"
    hdr_cells[3].text = "Coût Total"
    
    resources = [
        ("Serveur Cloud (VPS)", "1 mois", "50 DH", "50 DH"),
        ("Base de données managée", "1 mois", "0 DH", "0 DH (gratuit)"),
        ("Outils de développement", "1", "0 DH", "0 DH (open source)"),
        ("Hébergement du rapport", "1", "0 DH", "0 DH"),
    ]
    
    for element, qty, unit_cost, total in resources:
        row_cells = table.add_row().cells
        row_cells[0].text = element
        row_cells[1].text = qty
        row_cells[2].text = unit_cost
        row_cells[3].text = total
    
    doc.add_heading("7.2 Heures-Personnes", level=2)
    
    hours_table = doc.add_table(rows=1, cols=4)
    hours_table.style = 'Light Grid Accent 1'
    
    hdr_cells = hours_table.rows[0].cells
    hdr_cells[0].text = "Activité"
    hdr_cells[1].text = "Heures"
    hdr_cells[2].text = "Taux Horaire"
    hdr_cells[3].text = "Coût"
    
    hours = [
        ("Planification & Design", "40h", "0 DH/h", "0 DH"),
        ("Développement Backend", "120h", "0 DH/h", "0 DH"),
        ("Développement Frontend", "120h", "0 DH/h", "0 DH"),
        ("Testing & QA", "60h", "0 DH/h", "0 DH"),
        ("Documentation", "40h", "0 DH/h", "0 DH"),
    ]
    
    for activity, hours_num, rate, cost in hours:
        row_cells = hours_table.add_row().cells
        row_cells[0].text = activity
        row_cells[1].text = hours_num
        row_cells[2].text = rate
        row_cells[3].text = cost
    
    doc.add_paragraph("Total: 0 DH (Projet académique sans frais directs)")

def add_resources_section(doc):
    """Ajoute la section des ressources"""
    doc.add_heading("8. Ressources Techniques et Humaines", level=1)
    
    doc.add_heading("8.1 Équipe du Projet", level=2)
    doc.add_paragraph("Composée d'étudiants en projet de fin d'études travaillant à titre d'équipe pluridisciplinaire.")
    
    doc.add_heading("8.2 Technologies Utilisées", level=2)
    
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Light Grid Accent 1'
    
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = "Domaine"
    hdr_cells[1].text = "Technologie"
    hdr_cells[2].text = "Version"
    
    technologies = [
        ("Backend Framework", "Laravel", "12.x"),
        ("Language (Backend)", "PHP", "8.3+"),
        ("Frontend Framework", "Vue.js", "3.x"),
        ("Build Tool (Frontend)", "Vite", "Latest"),
        ("Authentication", "Laravel Sanctum", "Latest"),
        ("Database", "MySQL/SQLite", "8.0+/3.x"),
        ("API Design", "RESTful", "-"),
        ("Testing", "PHPUnit / Vitest", "Latest"),
    ]
    
    for domain, tech, version in technologies:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = domain
        row_cells[1].text = tech
        row_cells[2].text = version
    
    doc.add_heading("8.3 Environnements", level=2)
    environments = [
        ("Développement", "Local machines + Git repository"),
        ("Staging/Test", "Test server for QA"),
        ("Production", "VPS or cloud hosting")
    ]
    
    for env, description in environments:
        doc.add_paragraph(f"{env}: {description}", style='List Bullet')

def add_risks_section(doc):
    """Ajoute la section analyse des risques"""
    doc.add_heading("9. Analyse des Risques", level=1)
    
    doc.add_paragraph(
        "Identification et évaluation des risques pouvant affecter la réussite du projet."
    )
    
    risks_table = doc.add_table(rows=1, cols=6)
    risks_table.style = 'Light Grid Accent 1'
    
    hdr_cells = risks_table.rows[0].cells
    hdr_cells[0].text = "Risque"
    hdr_cells[1].text = "Probabilité"
    hdr_cells[2].text = "Impact"
    hdr_cells[3].text = "Criticité"
    hdr_cells[4].text = "Mitigation"
    hdr_cells[5].text = "Responsable"
    
    risks = [
        ("Délai de développement insuffisant", "Moyenne", "Haute", "Haute", 
         "Prioriser les fonctionnalités essentielles", "Chef de projet"),
        ("Problèmes d'intégration API-Frontend", "Moyenne", "Haute", "Haute",
         "Tests réguliers d'intégration", "Lead Dev"),
        ("Perte de données pendant le développement", "Basse", "Très Haute", "Moyenne",
         "Backup quotidien, Version Control", "DBA"),
        ("Changements de spécifications", "Moyenne", "Moyenne", "Moyenne",
         "Change management process", "Chef de projet"),
        ("Problèmes de performance", "Moyenne", "Moyenne", "Moyenne",
         "Profiling et optimisation régulière", "Tech Lead"),
        ("Absence d'équipe membre", "Basse", "Haute", "Moyenne",
         "Documentation complète", "Chef de projet"),
    ]
    
    for risk, prob, impact, criticite, mitigation, responsible in risks:
        row_cells = risks_table.add_row().cells
        row_cells[0].text = risk
        row_cells[1].text = prob
        row_cells[2].text = impact
        row_cells[3].text = criticite
        row_cells[4].text = mitigation
        row_cells[5].text = responsible

def add_communication_section(doc):
    """Ajoute la section communication"""
    doc.add_heading("10. Stratégie de Communication", level=1)
    
    doc.add_heading("10.1 Plan de Communication", level=2)
    
    comm_table = doc.add_table(rows=1, cols=5)
    comm_table.style = 'Light Grid Accent 1'
    
    hdr_cells = comm_table.rows[0].cells
    hdr_cells[0].text = "Type de Communication"
    hdr_cells[1].text = "Fréquence"
    hdr_cells[2].text = "Public"
    hdr_cells[3].text = "Moyen"
    hdr_cells[4].text = "Responsable"
    
    communications = [
        ("Réunion d'équipe", "Quotidienne", "Équipe", "Standup 15 min", "Chef de projet"),
        ("Point avec sponsor", "Hebdomadaire", "Sponsor", "Réunion vidéo", "Chef de projet"),
        ("Rapport d'avancement", "Hebdomadaire", "Stakeholders", "Email/Document", "Chef de projet"),
        ("Revue technique", "Bi-hebdomadaire", "Tech team", "Code Review", "Tech Lead"),
        ("Documentation", "Continue", "Tous", "Wiki/Docs", "Équipe"),
    ]
    
    for comm_type, freq, public, moyen, resp in communications:
        row_cells = comm_table.add_row().cells
        row_cells[0].text = comm_type
        row_cells[1].text = freq
        row_cells[2].text = public
        row_cells[3].text = moyen
        row_cells[4].text = resp

def add_conclusion(doc):
    """Ajoute la conclusion"""
    doc.add_page_break()
    doc.add_heading("Conclusion", level=1)
    
    doc.add_paragraph(
        "Ce Plan de Management de Projet définit le cadre et la structure du projet SchoolHub. "
        "Il établit des objectifs clairs, identifie les ressources nécessaires, et met en place "
        "des mécanismes de suivi et de contrôle pour assurer la réussite du projet."
    )
    
    doc.add_paragraph(
        "Le projet s'engage à respecter les délais, maintenir la qualité du code et de la documentation, "
        "et fournir une solution robuste et scalable pour la gestion d'établissements scolaires."
    )
    
    doc.add_heading("Signatures d'Approbation", level=2)
    
    approval_table = doc.add_table(rows=4, cols=2)
    approval_table.style = 'Light Grid Accent 1'
    
    approval_table.rows[0].cells[0].text = "Chef de Projet"
    approval_table.rows[0].cells[1].text = "Signature: _________________ Date: _____"
    
    approval_table.rows[1].cells[0].text = "Sponsor / Encadrant"
    approval_table.rows[1].cells[1].text = "Signature: _________________ Date: _____"
    
    approval_table.rows[2].cells[0].text = "Lead Technique"
    approval_table.rows[2].cells[1].text = "Signature: _________________ Date: _____"
    
    approval_table.rows[3].cells[0].text = "Approuvé le"
    approval_table.rows[3].cells[1].text = f": {datetime.now().strftime('%d/%m/%Y')}"

def main():
    """Fonction principale"""
    # Créer un nouveau document
    doc = Document()
    
    # Configurer les styles
    setup_styles(doc)
    
    # Ajouter les sections
    add_title_page(doc)
    add_table_of_contents(doc)
    add_context_section(doc)
    add_wbs_section(doc)
    add_stakeholders_section(doc)
    add_schedule_section(doc)
    add_indicators_section(doc)
    add_deliverables_section(doc)
    add_budget_section(doc)
    add_resources_section(doc)
    add_risks_section(doc)
    add_communication_section(doc)
    add_conclusion(doc)
    
    # Sauvegarder le document
    output_path = r"c:\Users\issam\Desktop\Moumens projects\SchoolHub\PMP_SchoolHub_2026.docx"
    doc.save(output_path)
    
    print("Document PMP cree avec succes!")
    print(f"Fichier: {output_path}")
    print(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

if __name__ == "__main__":
    main()
